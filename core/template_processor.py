#core/template_processor.py
import logging
import json
import re
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import time

from core.llama_runner import run_llama_prompt
from utils.file_utils import extract_text_from_file

logger = logging.getLogger(__name__)

# --- Helper functions for DOCX manipulation ---

def _clear_paragraph_content(paragraph: Paragraph):
    """Clears all runs from a DOCX paragraph."""
    for run in paragraph.runs:
        paragraph._element.remove(run._element)
    if not paragraph.runs:
        paragraph.add_run('')

def _add_text_to_paragraph_smart(paragraph: Paragraph, text: str):
    """
    Adds text to a DOCX paragraph. If text contains newlines, it creates new paragraphs
    or inserts line breaks. Attempts to handle simple bullet points.
    """
    _clear_paragraph_content(paragraph)
    
    lines = str(text).strip().split('\n')
    
    if not lines:
        return

    first_line = lines[0].strip()
    if first_line:
        if first_line.startswith(('* ', '- ')):
            paragraph.add_run('• ' + first_line[2:])
        else:
            paragraph.add_run(first_line)

    parent_element = paragraph._element.getparent()
    if parent_element is None:
        logger.warning("Could not find parent for paragraph. Cannot add subsequent lines as new paragraphs.")
        for line in lines[1:]:
            if line.strip():
                paragraph.add_run('\n')
                paragraph.add_run(line.strip())
            else:
                paragraph.add_run('\n')
        return

    for line_text in lines[1:]:
        if line_text.strip():
            new_p_element = parent_element.insert_paragraph_after(paragraph._element)
            new_p = Paragraph(new_p_element, paragraph.document)
            
            if line_text.strip().startswith(('* ', '- ')):
                new_p.add_run('• ' + line_text.strip()[2:])
            else:
                new_p.add_run(line_text.strip())
            paragraph = new_p
        else:
            new_p_element = parent_element.insert_paragraph_after(paragraph._element)
            paragraph = Paragraph(new_p_element, paragraph.document)


def _clear_cell_content(cell: _Cell):
    """Clears all content from a DOCX table cell."""
    for p in cell.paragraphs:
        cell._element.remove(p._element)
    if not cell.paragraphs:
        cell.add_paragraph()

def _add_text_to_cell_smart(cell: _Cell, text: str):
    """
    Adds text to a DOCX table cell, handling newlines as new paragraphs within the cell
    and attempting to format bullet points.
    """
    _clear_cell_content(cell)
    
    lines = str(text).strip().split('\n')
    
    if not lines:
        return

    target_paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    target_paragraph.clear()

    if lines and lines[0].strip():
        first_line = lines[0].strip()
        if first_line.startswith(('* ', '- ')):
            target_paragraph.add_run('• ' + first_line[2:])
        else:
            target_paragraph.add_run(first_line)
    
    for line in lines[1:]:
        if line.strip():
            new_paragraph = cell.add_paragraph()
            formatted_line = line.strip()
            if formatted_line.startswith(('* ', '- ')):
                new_paragraph.add_run('• ' + formatted_line[2:])
            else:
                new_paragraph.add_run(formatted_line)
        else:
            cell.add_paragraph("")


# --- Generic Dynamic Template Filling Logic ---

def fill_dynamic_template(template_path: str, resume_data: dict, output_path: str, template_schema: dict = None) -> dict:
    start_total_time = time.perf_counter()
    logger.info("Starting dynamic template filling for '%s'.", template_path)
    try:
        start_load_time = time.perf_counter()
        template_doc = Document(template_path)
        template_raw_text = extract_text_from_file(template_path)
        end_load_time = time.perf_counter()
        logger.info("Template document loaded and text extracted (Duration: %.4f s).", end_load_time - start_load_time)
    except Exception as e:
        logger.error("Failed to load template DOCX file: %s", e)
        return {"error": f"Failed to load template DOCX file: {e}"}

    if not resume_data:
        logger.error("No resume data provided. Cannot fill template dynamically.")
        return {"error": "No resume data provided for filling the template."}

    # Context for LLM: the raw template text, the structured resume data, AND the template_schema
    llm_context = f"EMPTY TEMPLATE CONTENT:\n{template_raw_text}\n\nRESUME DATA:\n{json.dumps(resume_data, indent=2)}"
    if template_schema:
        llm_context += f"\n\nTEMPLATE SCHEMA:\n{json.dumps(template_schema, indent=2)}"


    prompt = f"""
    You are an intelligent document automation assistant. Your task is to precisely fill the provided EMPTY TEMPLATE CONTENT using the RESUME DATA.
    If a TEMPLATE SCHEMA is provided, use it as a guide for the template's structure and expected fields.

    You MUST output a single, valid JSON object where keys are descriptive labels for template sections/fields, and values are the corresponding content from the RESUME DATA, formatted appropriately for a Word document.

    STRICT GUIDELINES FOR CONTENT AND FORMATTING:
    1.  **Strictly from Context:** All content provided in the values MUST be directly extracted or logically synthesized *only* from the provided RESUME DATA. DO NOT hallucinate, infer, or invent any information.
    2.  **Adhere to Template Schema (if provided):** If a TEMPLATE SCHEMA is present in the context, your JSON output keys should primarily align with the field names or section names described within that schema.
    3.  **Confidence Threshold:** Include a section/field in your JSON output *only* if you can confidently and accurately fill it with relevant information from the RESUME DATA. If a template section cannot be filled confidently, OMIT that key-value pair entirely.
    4.  **Repeating Sections (Advanced):** If the TEMPLATE SCHEMA indicates a section is "is_repeatable" (like multiple job experiences or education entries), your JSON value for that section should be a LIST of objects, where each object represents one instance of that repeatable block, following the `example_row_fields` or `fields` defined in the schema for that repeatable section. Fill as many instances as found in the RESUME DATA.
    5.  **Formatting:**
        * For lists (e.g., responsibilities, skills), use bullet points ('• ') or comma-separated lists, as appropriate. Use '\\n' for new lines between items or paragraphs.
        * For table cells, ensure content is concise.

    Example JSON output structure (ensure all values are strings for single fields, or lists of strings/objects for repeating content):
    ```json
    {{
      "Full Name": "{{resume_data.get('full_name', 'N/A')}}",
      "Contact Information": "Email: {{resume_data.get('email', 'N/A')}}\\nPhone: {{resume_data.get('phone_number', 'N/A')}}\\nLinkedIn: {{resume_data.get('linkedin_profile', 'N/A')}}",
      "Professional Summary Paragraph": "{{resume_data.get('professional_summary', 'N/A')}}",
      "Experience Section Details": [
        {{
          "Job Title": "Senior Software Engineer",
          "Company Name": "Tech Innovations",
          "Duration": "Jan 2020 - Present",
          "Responsibilities_Bullet_Points": "• Led development of scalable microservices.\\n• Mentored junior engineers."
        }},
        {{
          "Job Title": "Software Developer",
          "Company Name": "Web Solutions",
          "Duration": "Jul 2017 - Dec 2019",
          "Responsibilities_Bullet_Points": "• Built responsive UI components.\\n• Performed code reviews."
        }}
      ],
      "Skills List": "{{', '.join(resume_data.get('skills', []))}}"
    }}
    ```
    
    Generate the JSON object for filling the provided TEMPLATE CONTENT based on RESUME DATA and TEMPLATE SCHEMA.
    ```json
    """

    start_llm_time = time.perf_counter()
    llm_filling_instructions_raw = run_llama_prompt(prompt, context=llm_context, max_new_tokens=2000, temperature=0.1) 
    end_llm_time = time.perf_counter()
    logger.info("LLM generated filling instructions (Duration: %.4f s). Length: %d", end_llm_time - start_llm_time, len(llm_filling_instructions_raw))
    logger.debug("LLM Instructions raw: %s", llm_filling_instructions_raw[:500])

    try:
        start_json_parse_time = time.perf_counter()
        json_match = re.search(r'```json\n(.*)\n```', llm_filling_instructions_raw, re.DOTALL)
        if json_match:
            filling_instructions = json.loads(json_match.group(1).strip())
        else:
            filling_instructions = json.loads(llm_filling_instructions_raw.strip())
        end_json_parse_time = time.perf_counter()
        logger.info("Successfully parsed LLM filling instructions (Duration: %.4f s).", end_json_parse_time - start_json_parse_time)
        logger.debug("Parsed instructions: %s", json.dumps(filling_instructions, indent=2)[:500])
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse LLM filling instructions as JSON: {e}")
        logger.error(f"Malformed JSON from LLM:\n{llm_filling_instructions_raw}")
        return {"error": f"LLM did not provide valid JSON instructions for filling: {e}"}
    except Exception as e:
        logger.error(f"Unexpected error parsing LLM instructions: {e}", exc_info=True)
        return {"error": f"Failed to interpret LLM instructions: {e}"}

    start_apply_instructions_time = time.perf_counter()
    logger.info("Applying LLM instructions to template document dynamically.")
    
    filled_any_content = False
    filled_instruction_keys = set()

    # Apply instructions to paragraphs
    for p_idx, p in enumerate(template_doc.paragraphs):
        p_text_lower = p.text.strip().lower()
        for key, content in filling_instructions.items():
            key_lower = key.lower()
            if key in filled_instruction_keys or not isinstance(content, str):
                continue
            
            match = re.search(r'\b' + re.escape(key_lower.replace('_content', '').replace('_', ' ')) + r'\b', p_text_lower)
            if match:
                logger.info("Attempting to fill paragraph for instruction key: '%s' (p_idx: %d, text: '%s')", key, p_idx, p_text_lower[:50])
                try:
                    _add_text_to_paragraph_smart(p, str(content))
                    filled_any_content = True
                    filled_instruction_keys.add(key)
                    logger.debug("Dynamically filled paragraph for key '%s'.", key)
                    break
                except Exception as fill_e:
                    logger.error(f"Error filling paragraph for key '{key}': {fill_e}", exc_info=True)

    # Apply instructions to tables
    for t_idx, table in enumerate(template_doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_text_lower = cell.text.strip().lower()

                for key, content in filling_instructions.items():
                    if key in filled_instruction_keys:
                        continue
                    
                    key_lower = key.lower()
                    match = re.search(r'\b' + re.escape(key_lower.replace('_content', '').replace('_', ' ')) + r'\b', cell_text_lower)
                    if match and isinstance(content, str):
                        logger.info("Attempting to fill table cell for instruction key: '%s' (t_idx: %d, r_idx: %d, c_idx: %d, text: '%s')", key, t_idx, r_idx, c_idx, cell_text_lower[:50])
                        try:
                            _add_text_to_cell_smart(cell, str(content))
                            filled_any_content = True
                            filled_instruction_keys.add(key)
                            logger.debug("Dynamically filled table cell for key '%s'.", key)
                            break
                        except Exception as fill_e:
                            logger.error(f"Error filling table cell for key '{key}': {fill_e}", exc_info=True)
                    
                    elif isinstance(content, list) and 'sections' in template_schema and len(template_schema['sections']) > t_idx and 'table_structure' in template_schema['sections'][t_idx] and template_schema['sections'][t_idx]['table_structure'].get('is_repeatable', False):
                        logger.warning(f"Repeatable table content for key '{key}' identified, but dynamic row creation for complex tables is not fully implemented for all fields. Will attempt to fill first instance or skip subsequent.")
                        pass

    if not filled_any_content:
        logger.warning("No dynamic content was inserted based on LLM instructions. Template might be too unique or instructions not precise enough. Remaining instructions: %s", json.dumps({k: v for k,v in filling_instructions.items() if k not in filled_instruction_keys}))
        if not filled_any_content:
            template_doc.add_paragraph("\n--- Generated Content (No Specific Sections Found) ---")
            template_doc.add_paragraph(json.dumps({k: v for k,v in filling_instructions.items() if k not in filled_instruction_keys}, indent=2))
            logger.info("Appended remaining LLM generated content at the end of the document.")
    end_apply_instructions_time = time.perf_counter()
    logger.info("Applying LLM instructions to document completed (Duration: %.4f s).", end_apply_instructions_time - start_apply_instructions_time)

    start_save_time = time.perf_counter()
    try:
        template_doc.save(output_path)
        end_save_time = time.perf_counter()
        logger.info("Dynamically filled template saved to: %s (Duration: %.4f s)", output_path, end_save_time - start_save_time)
        return {"status": "success", "output_path": output_path}
    except Exception as e:
        logger.error("Error saving dynamically filled template to %s: %s", output_path, e, exc_info=True)
        return {"error": f"Failed to save dynamically filled template: {e}"}


def fill_skill_matrix_template(template_path: str, resume_data: dict, output_path: str) -> dict:
    start_total_time = time.perf_counter()
    logger.info("Starting to fill Skill-Qualification Matrix template (specific logic): %s", template_path)
    try:
        template_doc = Document(template_path)
    except Exception as e:
        logger.error("Failed to load template DOCX file: %s", e)
        return {"error": f"Failed to load template DOCX file: {e}"}

    if not resume_data:
        logger.error("No resume data provided. Cannot fill template.")
        return {"error": "No resume data provided for filling the template."}

    if not template_doc.tables:
        logger.error("Skill-Qualification Matrix template has no tables. Cannot fill.")
        return {"error": "Template is not a recognized Skill-Qualification Matrix format (no tables found)."}

    main_table = template_doc.tables[0]
    logger.info("Iterating through table rows to find qualifications in Skill Matrix.")
    start_table_iteration_time = time.perf_counter()
    
    start_row_idx = -1
    for r_idx, row in enumerate(main_table.rows):
        if len(row.cells) > 0:
            first_cell_text = row.cells[0].text.strip()
            if re.match(r'^\d+\.', first_cell_text):
                start_row_idx = r_idx
                logger.info(f"Dynamically detected qualification start at row index: {start_row_idx} (Content: '{first_cell_text}')")
                break
    
    if start_row_idx == -1:
        logger.warning("Could not dynamically find a row starting with a numbered item (e.g., '1.') in column 0. Defaulting to assuming first data row is index 3.")
        start_row_idx = 3

    for r_idx in range(start_row_idx, len(main_table.rows)):
        row = main_table.rows[r_idx]
        qualification_text = ""
        try:
            if len(row.cells) < 4:
                logger.warning(f"Skipping row {r_idx} due to insufficient columns ({len(row.cells)} < 4).")
                continue

            qualification_text = row.cells[1].text.strip()
            if not qualification_text:
                logger.debug("Skipping empty qualification row at index: %d", r_idx)
                continue

            logger.info("Processing qualification: '%s' (Row %d)", qualification_text, r_idx)
            llm_context = json.dumps(resume_data, indent=2)

            combined_prompt = f"""
            You are an expert at extracting and summarizing resume data for a skill matrix.
            Given the following structured resume data for {resume_data.get('full_name', 'the candidate')},
            and a specific qualification: '{qualification_text}'.

            Your task is to generate TWO distinct pieces of content in a single JSON object:
            1.  "vendor_response_content": A concise, bulleted list of relevant experience from the 'experience' section that directly addresses the specific qualification. Format this as a single string, using bullet points (e.g., • Point 1\\n• Point 2) for clarity.
            2.  "customer_references_content": A concise, bulleted list of companies, locations, and durations where the candidate performed work that is relevant to the qualification. Each item should be: • Company Name, Location (Start Date - End Date). Ensure this content is directly supported by the provided resume data.

            You MUST output a single, valid JSON object with exactly these two keys: "vendor_response_content" and "customer_references_content".
            If no relevant information is found, the value for a key should be "N/A".
            Do not include any conversational text, explanations, or markdown outside the JSON block.
            """
            
            combined_llm_response_raw = run_llama_prompt(
                combined_prompt, context=llm_context, max_new_tokens=2100, temperature=0.1
            )
            
            combined_response_data = {}
            try:
                json_match = re.search(r'```json\n(.*)\n```', combined_llm_response_raw, re.DOTALL)
                json_str = json_match.group(1).strip() if json_match else combined_llm_response_raw.strip()
                combined_response_data = json.loads(json_str)
            except json.JSONDecodeError as e:
                logger.error(f"JSON Decode Error parsing combined LLM response for qualification '{qualification_text}': {e}")
                logger.error(f"Malformed Combined JSON from LLM:\n{combined_llm_response_raw}")
                _add_text_to_cell_smart(row.cells[2], "Error: Could not parse AI response.")
                _add_text_to_cell_smart(row.cells[3], "Error: Could not parse AI response.")
                continue
            
            vendor_content = combined_response_data.get('vendor_response_content', 'N/A')
            customer_content = combined_response_data.get('customer_references_content', 'N/A')

            # --- FIX START ---
            # This block ensures that if the LLM returns a list instead of a string,
            # we convert it to a string before trying to write it to the document. This prevents the crash.
            if isinstance(vendor_content, list):
                vendor_content = "\n".join(map(str, vendor_content))
            if isinstance(customer_content, list):
                customer_content = "\n".join(map(str, customer_content))
            # --- FIX END ---

            _add_text_to_cell_smart(row.cells[2], str(vendor_content))
            _add_text_to_cell_smart(row.cells[3], str(customer_content))

        except Exception as e:
            logger.error(f"Error processing row {r_idx} for qualification '{qualification_text}': {e}", exc_info=True)
            if len(row.cells) > 3:
                _add_text_to_cell_smart(row.cells[2], "Error during processing.")
                _add_text_to_cell_smart(row.cells[3], "Error during processing.")

    end_table_iteration_time = time.perf_counter()
    logger.info("Table iteration and LLM calls for Skill Matrix completed (Duration: %.4f s).", end_table_iteration_time - start_table_iteration_time)

    start_save_time = time.perf_counter()
    try:
        template_doc.save(output_path)
        end_save_time = time.perf_counter()
        logger.info("Dynamically filled template saved to: %s (Duration: %.4f s)", output_path, end_save_time - start_save_time)
        return {"status": "success", "output_path": output_path}
    except Exception as e:
        logger.error("Error saving dynamically filled template to %s: %s", output_path, e, exc_info=True)
        return {"error": f"Failed to save filled template: {e}"}