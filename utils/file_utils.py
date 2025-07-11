import logging
import os
from docx import Document
import mammoth
import fitz
import re
import time

logger = logging.getLogger(__name__)

def clean_text_with_regex(text: str) -> str:
    """
    Cleans text by removing non-printable ASCII characters, excessive whitespace,
    and standardizing newlines.
    """
    # Remove non-printable ASCII characters (except common whitespace: tab, newline, carriage return)
    # This regex matches any character that is NOT a printable ASCII character (0x20 to 0x7E)
    # and is NOT a tab (\t), newline (\n), or carriage return (\r).
    cleaned_text = re.sub(r'[^\x20-\x7E\t\n\r]', '', text)
    
    # Replace multiple newlines with a maximum of two newlines
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    
    # Replace multiple spaces/tabs with a single space
    cleaned_text = re.sub(r'[ \t]+', ' ', cleaned_text)
    
    # Strip leading/trailing whitespace from each line and the overall text
    cleaned_text = '\n'.join([line.strip() for line in cleaned_text.split('\n')])
    cleaned_text = cleaned_text.strip()
    
    return cleaned_text


def extract_text_from_file(file_path: str) -> str:
    """
    Extracts text from various document types (.pdf, .doc, .docx) and cleans it.
    Args:
        file_path (str): The path to the file.
    Returns:
        str: The extracted and cleaned text content of the file.
    Raises:
        ValueError: If the file format is unsupported.
        Exception: If an error occurs during text extraction.
    """
    start_time = time.perf_counter()
    ext = os.path.splitext(file_path)[1].lower()
    logger.info("Starting text extraction for file: %s with extension '%s'", file_path, ext)

    text = "" # Initialize text here

    if ext == ".pdf":
        logger.info("Detected format: PDF. Attempting to extract text using PyMuPDF.")
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except Exception as e:
            logger.error("Failed to extract text from PDF '%s': %s", file_path, e, exc_info=True)
            raise
    elif ext == ".docx":
        logger.info("Detected format: DOCX. Attempting to extract text using python-docx.")
        text_content = []
        try:
            doc = Document(file_path)
            
            # Extract text from top-level paragraphs
            for para in doc.paragraphs:
                text_content.append(para.text)
            
            # Extract text from tables (NEW LOGIC)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs: # Iterate paragraphs within each cell
                            text_content.append(paragraph.text)

            text = "\n".join(text_content) # Join all collected text
        except Exception as e:
            logger.error("Failed to extract text from DOCX '%s': %s", file_path, e, exc_info=True)
            raise
    elif ext == ".doc":
        logger.info("Detected format: DOC. Attempting to extract text using mammoth.")
        try:
            with open(file_path, "rb") as f:
                result = mammoth.extract_raw_text(f)
            text = result.value
        except Exception as e:
            logger.error("Failed to extract text from DOC '%s': %s", file_path, e, exc_info=True)
            raise
    else:
        logger.error("Unsupported file format encountered: '%s' for file '%s'.", ext, file_path)
        raise ValueError(f"Unsupported file format: {ext}")

    # Apply cleaning after text extraction
    if text: # Only clean if text was successfully extracted
        cleaned_text = clean_text_with_regex(text)
        end_time = time.perf_counter()
        logger.info("Successfully extracted and cleaned text from file. Total original length: %d, Cleaned length: %d (Duration: %.4f s)", len(text), len(cleaned_text), end_time - start_time)
        return cleaned_text
    else:
        end_time = time.perf_counter()
        logger.warning("No text extracted or text was empty after extraction for file: %s (Duration: %.4f s)", file_path, end_time - start_time)
        return ""


def chunk_text(text: str, chunk_size_words: int = 200, chunk_overlap_words: int = 50) -> list[str]:
    """
    Splits text into smaller, overlapping chunks suitable for embedding.
    Chunks are created based on word count for simplicity.
    """
    start_time = time.perf_counter()
    logger.info("Chunking text (approx. words: %d) with chunk_size %d, overlap %d.",
                len(text.split()), chunk_size_words, chunk_overlap_words)
    
    words = text.split()
    chunks = []
    
    i = 0
    while i < len(words):
        chunk = words[i : i + chunk_size_words]
        chunks.append(" ".join(chunk))
        
        i += (chunk_size_words - chunk_overlap_words)
        if i >= len(words):
            break
    end_time = time.perf_counter()
    logger.info("Text chunked into %d chunks (Duration: %.4f s).", len(chunks), end_time - start_time)
    return chunks

def get_combined_context(resume_path: str, template_path: str) -> str:
    """
    Generates a combined text context from a resume and a template file.
    (Primarily for logging/debugging in a RAG setup, as LLM context comes from ChromaDB)
    """
    logger.debug("Generating combined raw text context.")
    resume_text = ""
    try:
        resume_text = extract_text_from_file(resume_path)
    except Exception as e:
        logger.error("Could not extract text from resume '%s' for combined context: %s", resume_path, e)

    template_text = ""
    try:
        template_text = extract_text_from_file(template_path)
    except Exception as e:
        logger.error("Could not extract text from template '%s' for combined context: %s", template_path, e)

    combined = f"RESUME:\n{resume_text.strip()}\n\nTEMPLATE:\n{template_text.strip()}"
    return combined